import docx
import pandas as pd
import numpy as np
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


class WriteWord(object):

    """
    利用python-docx在Word写入

    # 主要参考链接
    https://python-docx.readthedocs.io/en/latest/
    https://python-docx.readthedocs.io/en/latest/_modules/docx/document.html

     https://blog.csdn.net/xtfge0915/article/details/83478639
    表格 https://blog.csdn.net/xtfge0915/article/details/83479933
    表格样式 https://blog.csdn.net/xtfge0915/article/details/83480120
    """

    def __init__(self):

        self.doc = docx.Document()
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    def add_paragraph(self, text, blank=False, align=False, line_spacing=True,
                      size=10, bold=False, italic=False):

        """ 插入段落 """

        para = self.doc.add_paragraph()
        if blank:
            text = "    " + text
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 必须加上此行 否则样式无效

        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor(0, 0, 0)  # 段落字体颜色

        if line_spacing:
            para.paragraph_format.line_spacing = Pt(int(1.1*size))
            para.paragraph_format.space_before = Pt(int(1.1*size))  # 段落前间距
            para.paragraph_format.space_after = Pt(int(1.1*size))  # 段落后间距
        if align:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中

        # para.paragraph_format.line_spacing = Pt(18)  # 段落行间距
        # para.paragraph_format.left_indent = Inches(0.75) # 左缩进

    def add_picture(self, pic_file, width=5.0, height=None, align=True):

        """ 插入图片（图片是插入在段落里面的） """

        para = self.doc.add_paragraph()
        run = para.add_run()
        width = Inches(width)
        run.add_picture(pic_file, width=width, height=height)

        if align:
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中

    def add_picture_note(self, text, size=10, align=True):

        """ 图片备注 """

        self.add_paragraph(text, align=align, line_spacing=False,
                           size=size, bold=False, italic=False)

    def reshape_dataframe_columns(self, data):

        """ 将 dataframe columns 放到第一行 """

        index = list(data.index)
        data.loc['col', :] = data.columns
        index.insert(0, 'col')
        data = data.loc[index, :]
        return data

    def add_table_note(self, text, size=10, align=True):

        """ 表格备注 """

        self.add_paragraph(text, align=align, line_spacing=False,
                           size=size, bold=False, italic=False)

    def add_table(self, data=None, style="Table Grid", rows=2, cols=3,
                  height_list=None, width_list=None,
                  font_color=RGBColor(0, 0, 0), font_size=8, font_name="微软雅黑",
                  bold=False, italic=False, sepcial_header=True, format_force=None, align=True):

        """ 插入表格 """

        if data is not None:
            rows = data.shape[0]
            cols = data.shape[1]
        if height_list is None:
            height_list = rows * [0.45]
        if width_list is None:
            width_list = cols * [1.5]

        if format_force is None:
            format_force = pd.DataFrame([], index=data.index, columns=data.columns)

        format_force = format_force.fillna(True)
        table = self.doc.add_table(rows=rows, cols=cols, style=style)

        if data is not None:
            for i_row in range(rows):
                for i_col in range(cols):

                    table_row = table.rows[i_row]
                    cell = table_row.cells[i_col]
                    val = data.iloc[i_row, i_col]
                    # print(data.index[i_row], data.columns[i_col], type(val) == np.float)

                    if format_force.iloc[i_row, i_col]:
                        try:
                            val = np.float(val)
                        except Exception as e:
                            pass

                    if format_force.iloc[i_row, i_col] and type(val) == np.float:
                        val = '{:.2%}'.format(val)

                    run = cell.paragraphs[0].add_run(str(val))

                    # 表格文字对齐样式
                    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    # 表格文字字体
                    run.font.name = font_name
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = font_color
                    run.font.bold = bold
                    run.font.italic = italic

                    # 表格长宽
                    cell.width = Cm(width_list[i_col])

                    # 对表头特殊处理
                    if sepcial_header and i_row == 0:
                        run.font.bold = True
                        table_row.height = Cm(1.3*height_list[i_row])
                        # print(table_row.height, i_row)
                    else:
                        table_row.height = Cm(height_list[i_row])
                        # print(table_row.height, i_row)

        # 表格整体样式
        table.autofit = True
        if align:
            table.alignment = WD_TAB_ALIGNMENT.CENTER

        # table.width = Inches(10.0)
        # table.height = Inches(5.0)

    def add_footer(self, text):

        """ 插入页脚 """
        # 有些问题 暂时利用win32word的插入页脚功能
        section = self.doc.sections[0]
        footers = section.footer  # a HeadersFooters collection object
        footers.text = text
        footers.is_linked_to_previous = False

    def add_page_break(self):

        """ 插入页面分隔 """

        self.doc.add_page_break()

    def save(self, filename):

        """ 保存文件 """

        self.doc.save(filename)

    def demo(self):

        """ 举例 """
        pic_file = r'E:\Data\fund_data\fund_index_exposure_weekly\output_exposure\half_year_date.png'
        file = r"C:\Users\doufucheng\OneDrive\Desktop\demo.docx"

        self.add_paragraph("题目1", size=24, align=True, line_spacing=False, bold=True)
        self.add_paragraph("标题1", size=18, bold=True)
        self.add_paragraph(5 * "这里是第一个段落，飞流直下三千尺，她在丛中笑。", blank=True)
        self.add_paragraph(5 * "这里是第二个段落，杨花落尽子规啼，春风吹又生。", blank=True)
        self.add_paragraph("标题2", size=18, bold=True)
        self.add_paragraph(5 * "这里是第三个段落，只恐双溪舴艋舟，载不动许多愁。", blank=True)
        self.add_picture(pic_file=pic_file)
        self.add_picture_note("图1：基金风格暴露20190101")
        data = pd.DataFrame([["日期", "20190101"],
                             ['3.7', 0.8]],
                            index=['行索引1', '行索引2'],
                            columns=['列索引1', '列索引2'])
        data = self.reshape_dataframe_columns(data)
        # style="Medium Grid 1 Accent 1"
        self.add_table(data=data, width_list=[2.4, 3.6], height_list=[1.0, 1.0, 0.5])
        self.add_footer(5*"页脚")
        self.save(file)


if __name__ == '__main__':

    self = WriteWord()
    self.demo()
